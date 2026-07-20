#!/usr/bin/env python3
"""DOCX 后处理：表格实线边框、微软雅黑、居中对齐、页码页脚"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = tcBorders.find(qn(f'w:{edge}'))
            if border is None:
                border = parse_xml(f'<w:{edge} {nsdecls("w")}></w:{edge}>')
                tcBorders.append(border)
            border.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            border.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            border.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            border.set(qn('w:space'), '0')


def process_docx(input_path, output_path):
    doc = Document(input_path)

    # 表格处理
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')

        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.25
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(10.5)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

                if ri == 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                set_cell_border(cell,
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'},
                )

    # 页码页脚
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run()
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fldChar2)

    # 全文微软雅黑
    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        text = p.text.strip() if p.text else ''
        if text.startswith('图') or text.startswith('表'):
            is_italic = any(run.font.italic for run in p.runs)
            if is_italic:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(output_path)
    print(f'✅ DOCX saved: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: post_process_docx.py <input.docx> <output.docx>')
        sys.exit(1)
    process_docx(sys.argv[1], sys.argv[2])
