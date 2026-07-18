#!/usr/bin/env python3
"""
后处理文档.py <输入.docx> <输出.docx>

全文档微软雅黑 + 表格实线边框 + 页码页脚 + 封面识别格式化。

封面识别逻辑：
  查找第一个分页符（来自 Markdown 的 \\newpage），
  之前的所有段落视为封面区域，应用居中+大字。

用法：
  python3 脚本/后处理文档.py <输入.docx> <输出.docx>
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = tcBorders.find(qn(f'w:{edge}'))
            if border is None:
                border = parse_xml(f'<w:{edge} {nsdecls("w")}></w:{edge}>')
                tcBorders.append(border)
            border.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            border.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            border.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            border.set(qn('w:space'), '0')


def set_run_font(run, name='Microsoft YaHei', size=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    rpr = run._element.rPr
    if rpr is None:
        run._element.rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        rpr = run._element.rPr
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:eastAsia'), name)


def find_first_page_break(doc):
    """
    查找第一个分页符段落索引。
    新页过滤.lua 插入 <w:br w:type="page"/>，对应 python-docx 中的
    run._element.xml 包含 'w:br' + 'w:type=\"page\"'。
    返回段落索引（首个分页符之后的段落），若无则返回 None。
    """
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            xml = run._element.xml
            if 'w:br' in xml and 'page' in xml:
                return i + 1  # 封面到分页符之前
    return None


def format_cover(paragraphs):
    """将段落列表格式化为封面。"""
    if not paragraphs:
        return

    for idx, p in enumerate(paragraphs):
        text = p.text.strip() if p.text else ''
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Determine font size based on position
        if idx == 0 and len(text) > 3:
            # 标题：三号(16pt) 加粗
            fs = 16
            bold = True
            p.paragraph_format.space_after = Pt(6)
        elif idx == 1 and len(text) < 20:
            # 作者行：14pt
            fs = 14
            bold = False
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        else:
            # 版本/描述等：12pt
            fs = 12
            bold = False

        for run in p.runs:
            set_run_font(run, size=fs)
            run.font.bold = bold


def process_docx(input_path, output_path):
    doc = Document(input_path)

    # ---- 0. 查找封面和分页符 ----
    first_body_idx = find_first_page_break(doc)

    # ---- 1. Table formatting (unchanged) ----
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')

        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.25
                    for run in p.runs:
                        set_run_font(run, size=10.5)
                if ri == 0:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.bold = True
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_cell_border(cell,
                    top={'val': 'single', 'sz': '4', 'color': '000000'},
                    bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                    left={'val': 'single', 'sz': '4', 'color': '000000'},
                    right={'val': 'single', 'sz': '4', 'color': '000000'})

    # ---- 2. Page footer (unchanged) ----
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run, size=9)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fldChar2)

    # ---- 3. Cover page formatting ----
    if first_body_idx and first_body_idx > 1:
        cover_paras = doc.paragraphs[:first_body_idx]
        format_cover(cover_paras)
    else:
        # 无分页符时，按内容模式猜测封面：前1~4段含标题的当封面
        cover_count = 0
        for p in doc.paragraphs[:6]:
            t = p.text.strip() if p.text else ''
            if t and cover_count < 4:
                cover_count += 1
            elif not t:
                break
        if cover_count >= 1:
            format_cover(doc.paragraphs[:cover_count])

    # ---- 4. Body font & alignment ----
    body_start = first_body_idx if first_body_idx else 0
    for p in doc.paragraphs[body_start:]:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        # 左对齐（正文段落）
        if not p.style.name.startswith('Heading'):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, size=12)

    # ---- 5. 插图标题居中（覆盖上一步的左对齐） ----
    for p in doc.paragraphs[body_start:]:
        text = p.text.strip() if p.text else ''
        if (text.startswith('图') or text.startswith('表')):
            # 检测是否为 italic（插图标题通常 pandoc 渲染为斜体）
            is_italic = any(run.font.italic for run in p.runs)
            if is_italic:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"✅ DOCX post-processed: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: post_process_docx.py <input.docx> <output.docx>")
        sys.exit(1)
    process_docx(sys.argv[1], sys.argv[2])
